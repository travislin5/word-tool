import os
from typing import Optional
from xml.dom.minidom import Document
from fastapi import FastAPI, Form, HTTPException, Path, UploadFile
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from docx import Document


app = FastAPI()


class CompanyRequest(BaseModel):
    company_name: str


class Item(BaseModel):
    file: Optional[UploadFile] = None
    key1: Optional[str] = None
    val1: Optional[str] = None
    key2: Optional[str] = None
    val2: Optional[str] = None


# 將 /image 資料夾掛載到 /static 路徑
app.mount("/image", StaticFiles(directory="image"), name="image")


# 首頁設定成某個html
@app.get("/", response_class=FileResponse)
async def read_root():
    return FileResponse("./public/index.html")


@app.post("/get_company_data")
def get_company_data(request: CompanyRequest):
    test1()
    print("********")


@app.post("/test")
async def test333(
    file: Optional[UploadFile] = Form(None),
    key1: Optional[str] = Form(None),
    val1: Optional[str] = Form(None),
    key2: Optional[str] = Form(None),
    val2: Optional[str] = Form(None),
):
    item = Item(
        file=file,
        key1=key1,
        val1=val1,
        key2=key2,
        val2=val2,
    )
    print(item)

    extension = os.path.splitext(item.file.filename)[1].lower()
    file_n = os.path.splitext(item.file.filename)[0].lower()

    if extension not in [".doc", ".docx"]:
        raise HTTPException(
            status_code=400,
            detail="Invalid file extension. Only .doc and .docx files are allowed.",
        )
    file_location = "image/" + file_n + "_afterFix" + ".docx"  # 文件保存位置
    with open(file_location, "wb") as f:
        f.write(await file.read())  # 寫入文件

    doc = Document(file_location)

    if item.key1 is not None:
        for para in doc.paragraphs:
            if item.key1 in para.text:
                for run in para.runs:
                    if item.key1 in run.text:
                        run.text = run.text.replace(item.key1, item.val1)

    if item.key2 is not None:
        for para in doc.paragraphs:
            if item.key2 in para.text:
                for run in para.runs:
                    if item.key2 in run.text:
                        run.text = run.text.replace(item.key2, item.val2)

    doc.save("image/" + file_n + "_afterFix" + ".docx")
    word_path = "image/" + file_n + "_afterFix" + ".docx"
    return {"screenshot_path": f"{word_path}"}
